from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import io
import os


def add_bottom_border(paragraph):
    """Add a bottom border to a paragraph"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    # Create a bottom border element
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')      # Line thickness
    bottom.set(qn('w:space'), '1')   # Padding
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    
    pPr.append(pBdr)

def get_template_path():
    """Get the template file path, trying multiple locations"""
    current_dir = os.path.dirname(os.path.abspath(__file__))
    possible_paths = [
        os.path.join(current_dir, '../templates/sample_formatted.docx'),
        os.path.join(current_dir, 'templates/sample_formatted.docx'),
        'templates/sample_formatted.docx',
        os.path.join(os.path.dirname(current_dir), 'templates/sample_formatted.docx')
    ]
    for path in possible_paths:
        if os.path.exists(path):
            return path
    raise FileNotFoundError("Template file not found. Please ensure templates/sample_formatted.docx exists.")

def split_single_line_transcript(text: str) -> list:
    """
    Split a single-line transcript into logical segments.
    Detects speakers (names ending with ": ") and music markers.
    """
    lines = []
    segments = []
    last_pos = 0
    
    # Find all speaker segments
    for match in re.finditer(r'[A-Z][A-Za-z. ]+: ', text):
        start = match.start()
        
        # Add anything before this speaker as a segment
        if start > last_pos:
            before_text = text[last_pos:start].strip()
            if before_text and not before_text.startswith('♪'):
                segments.append(('narr', before_text))
        
        # Find the next speaker or end of text
        next_speaker = re.search(r'[A-Z][A-Za-z. ]+: ', text[start + match.end():])
        if next_speaker:
            # Speaker segment
            speaker_end = start + match.end() + next_speaker.start()
            segments.append(('speaker', text[start:speaker_end].strip()))
            last_pos = speaker_end
        else:
            # Last speaker segment to end
            segments.append(('speaker', text[start:].strip()))
            last_pos = len(text)
            break
    
    # Handle remaining text
    if last_pos < len(text):
        remaining = text[last_pos:].strip()
        if remaining:
            segments.append(('narr', remaining))
    
    # Convert segments to lines
    for seg_type, content in segments:
        lines.append(content)
    
    return lines if lines else [text]  # Fallback to original if no segments found

def apply_text_formatting(text: str, para, add_newline: bool = False):
    """
    Apply professional formatting to text:
    - Italicize show names (World Impact)
    - Bold websites/organizations (worldimpact.tv, ORU)
    - Bold scripture references (1 John 2:18)
    - Handle song titles with italics and quotes
    """
    # Regular expressions for different formatting patterns
    patterns = [
        # Scripture references in format "1 John 2:18" or "2 Timothy 3:1-5"
        (r'(\d+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\s+\d+(?:-\d+)?:\d+(?:-\d+)?)', 'BOLD'),
        
        # Show names (World Impact, case-insensitive)
        (r'(\bWorld Impact\b)', 'ITALIC'),
        
        # URLs and websites
        (r'(\bworldimpact\.tv\b)', 'BOLD'),
        
        # Organization names (ORU, etc.)
        (r'(\bORU\b)', 'BOLD'),
        (r'(\bOral Roberts University\b)', 'BOLD'),
        
        # Song titles in quotes
        (r'("Give Me Jesus")', 'ITALIC'),
    ]
    
    current_pos = 0
    text_length = len(text)
    
    while current_pos < text_length:
        # Find the nearest pattern match
        next_match = None
        next_pos = text_length
        match_type = None
        
        for pattern, fmt_type in patterns:
            pattern_obj = re.compile(pattern, re.IGNORECASE)
            match = pattern_obj.search(text, current_pos)
            if match and match.start() < next_pos:
                next_pos = match.start()
                next_match = match
                match_type = fmt_type
        
        # Add text before the next match
        if next_pos > current_pos:
            normal_text = text[current_pos:next_pos]
            if normal_text:
                run = para.add_run(normal_text)
                run.font.name = 'Calibri'
                run.font.size = Pt(12)
                run.bold = False
                run.italic = False
        
        # Add formatted match
        if next_match:
            matched_text = next_match.group(0)
            current_pos = next_match.end()
            
            run = para.add_run(matched_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            
            if match_type == 'BOLD':
                run.bold = True
                run.italic = False
            elif match_type == 'ITALIC':
                run.italic = True
                run.bold = False
            else:
                run.bold = False
                run.italic = False
        else:
            break
    
    # Add remaining text
    if current_pos < text_length:
        remaining = text[current_pos:]
        if remaining:
            run = para.add_run(remaining)
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.bold = False
            run.italic = False

def format_transcript(text: str, title: str = None) -> bytes:
    """
    ORU Professional Transcript Formatter
    Formats raw transcript following ORU standards using rule-based parsing.
    Uses existing template with branding and appends formatted transcript content.
    
    Formatting rules:
    - Base: Calibri 12pt, 1.15x line spacing
    - Title: Calibri 18pt bold, centered, border beneath, 12pt before/18pt after
    - Speakers: Bold name, normal content on same line, 6pt after
    - Music: Centered, italicized, 3pt after
    - Scripture: Italic, 0.5" indent, 6pt after
    - Narration: 0.25" first-line indent, 6pt after
    - No punctuation at line starts
    """
    # Get template path and open the template
    template_path = get_template_path()
    doc = Document(template_path)
    
    # If the text is all on one line (no \n), split it intelligently
    if '\n' not in text.strip():
        # Single-line input: split by speaker patterns and music markers
        lines = split_single_line_transcript(text)
    else:
        lines = text.split('\n')
    
    # Track state
    title_added = False
    
    # Process each line from the input text
    for line in lines:
        line = line.strip()
        
        # Skip empty lines initially, then add appropriate spacing
        if not line:
            if title_added:
                doc.add_paragraph()  # Add blank line
            continue
        
        # Detect title (first meaningful line)
        if not title_added:
            title_para = doc.add_paragraph(line)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Format title: Calibri 18pt bold
            for run in title_para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(18)
                run.bold = True
            
            # Set spacing: 12 pt before, 18 pt after
            title_para.paragraph_format.space_before = Pt(12)
            title_para.paragraph_format.space_after = Pt(18)
            
            # Add bottom border
            add_bottom_border(title_para)
            
            title_added = True
            continue
        
        # Detect speaker lines (name ending with colon)
        # Split speaker name from content if both present
        if ':' in line:
            parts = line.split(':', 1)
            speaker_name = parts[0].strip()
            content = parts[1].strip() if len(parts) > 1 else ""
            
            # Check if it looks like a speaker pattern
            # Accept if: short name (<60 chars) OR contains title keywords OR looks like a proper name
            is_speaker = (
                len(speaker_name) < 60 or 
                any(keyword in speaker_name.lower() for keyword in ['dr.', 'mrs.', 'mr.', 'ms.', 'prof.', 'pastor', 'announcer']) or
                (len(speaker_name) < 100 and speaker_name[0].isupper() and len(speaker_name.split()) <= 3)
            )
            
            # Don't treat scripture citations as speakers
            if not (speaker_name.lower().startswith('mark') and 'this' in speaker_name.lower()):
                if is_speaker:
                    # Speaker line
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    para.paragraph_format.line_spacing = 1.15
                    
                    # Add speaker name in bold
                    speaker_run = para.add_run(speaker_name + ': ')
                    speaker_run.font.name = 'Calibri'
                    speaker_run.font.size = Pt(12)
                    speaker_run.bold = True
                    speaker_run.italic = False
                    
                    # Add content with professional formatting
                    if content:
                        apply_text_formatting(content, para)
                    
                    para.paragraph_format.space_after = Pt(6)
                    continue
            
            # Detect music lines
        if '♪' in line:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.line_spacing = 1.15
            
            # Add music line with italic formatting
            music_run = para.add_run(line)
            music_run.font.name = 'Calibri'
            music_run.font.size = Pt(12)
            music_run.italic = True
            music_run.bold = False
            
            para.paragraph_format.space_after = Pt(3)
            continue
        
        # Detect scripture references (pattern: number Book chapter:verse)
        scripture_pattern = re.compile(r'^\d+\s+[A-Z][a-z]+\s+\d+:\d+', re.IGNORECASE)
        if scripture_pattern.match(line) or re.search(r'\d+\s+[A-Z][a-z]+\s+\d+:\d+', line):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.line_spacing = 1.15
            para.paragraph_format.left_indent = Inches(0.5)
            
            # Scripture references should be in bold
            scripture_run = para.add_run(line)
            scripture_run.font.name = 'Calibri'
            scripture_run.font.size = Pt(12)
            scripture_run.bold = True
            scripture_run.italic = False
            
            para.paragraph_format.space_after = Pt(6)
            continue
        
        # Regular narration/descriptive text
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para.paragraph_format.line_spacing = 1.15
        para.paragraph_format.first_line_indent = Inches(0.25)
        
        # Apply professional formatting to the text
        apply_text_formatting(line, para)
        
        para.paragraph_format.space_after = Pt(6)
    
    # Note: Footer is already set up in the template with branding
    # No need to add or modify it
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.read()

