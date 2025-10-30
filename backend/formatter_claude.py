from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import os
from anthropic import Anthropic
from dotenv import load_dotenv

# Load environment variables
load_dotenv()


def add_bottom_border(paragraph):
    """Add a bottom border to a paragraph"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    
    pPr.append(pBdr)


def get_template_path():
    """Get the template file path"""
    current_dir = os.path.dirname(os.path.abspath(__file__))
    possible_paths = [
        os.path.join(current_dir, '../templates/sample_formatted.docx'),
        os.path.join(current_dir, 'templates/sample_formatted.docx'),
        'templates/sample_formatted.docx',
        os.path.join(os.path.dirname(current_dir), 'templates/sample_formatted.docx')
    ]
    for path in possible_paths:
        if os.path.exists(path):
            return path
    raise FileNotFoundError("Template file not found.")


def get_claude_client():
    """Initialize and return Claude client"""
    api_key = os.getenv('ANTHROPIC_API_KEY')
    if not api_key:
        raise ValueError("ANTHROPIC_API_KEY not found in environment")
    return Anthropic(api_key=api_key, timeout=120, max_retries=5)


def format_transcript_with_claude(text: str) -> bytes:
    """
    Use Claude AI to format the transcript according to professional standards.
    """
    client = get_claude_client()
    
    # Open the template
    template_path = get_template_path()
    doc = Document(template_path)
    
    # Prepare the comprehensive formatting instructions
    system_prompt = """You are a professional transcript formatting assistant specializing in converting raw TV show transcripts into polished, publication-ready documents.

Your task is to transform unformatted transcript text into structured content that I will then format into a Word document.

Return the formatted content as JSON with this structure:
{
  "title": "Episode Title",
  "segments": [
    {
      "type": "speaker|music|scripture|narration",
      "speaker": "Speaker name or null",
      "content": "Formatted content text",
      "emphasis": [
        {"text": "worldimpact.tv", "style": "bold"},
        {"text": "World Impact", "style": "italic"},
        {"text": "1 John 2:18", "style": "bold"},
        {"text": "Give Me Jesus", "style": "italic_quote"}
      ]
    }
  ]
}

FORMATTING RULES:
- Speaker names: Bold with colon (Dr. Billy Wilson:)
- Show names: Italic (World Impact)
- Websites: Bold (worldimpact.tv)
- Organizations: Bold (ORU, Oral Roberts University)
- Scripture references: Bold (1 John 2:18, 2 Timothy 3:1-5)
- Song titles: Italic with quotes ("Give Me Jesus")
- Music/lyrics: Preserve ♪ symbols, italic text
- Quoted scripture: Italic with quotes ("But mark this...")

Transform the raw transcript following these guidelines."""
    
    user_message = f"""Format this raw transcript according to professional standards:\n\n{text}"""
    
    # Call Claude
    response = client.messages.create(
        model="claude-sonnet-4-5-20250929",
        max_tokens=1200,
        temperature=0.0,
        top_p=1.0,
        system=system_prompt,
        messages=[
            {
                "role": "user",
                "content": user_message
            }
        ],
        stream=False
    )
    
    # Get the response text
    full_response = response.content[0].text if response.content else ""
    
    # Parse the JSON response
    # Extract JSON from markdown fences if present
    if '```json' in full_response:
        full_response = full_response.split('```json')[1].split('```')[0]
    elif '```' in full_response:
        full_response = full_response.split('```')[1]
    
    import json
    try:
        data = json.loads(full_response)
    except json.JSONDecodeError:
        # Fallback: create simple structure
        data = {"title": "World Impact Transcript", "segments": [{"type": "narration", "content": text}]}
    
    # Apply formatting to document
    if data.get("title"):
        title_para = doc.add_paragraph(data["title"])
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(18)
            run.bold = True
        title_para.paragraph_format.space_before = Pt(12)
        title_para.paragraph_format.space_after = Pt(18)
        add_bottom_border(title_para)
    
    # Process segments
    for segment in data.get("segments", []):
        seg_type = segment.get("type", "narration")
        content = segment.get("content", "")
        speaker = segment.get("speaker")
        emphasis = segment.get("emphasis", [])
        
        if seg_type == "speaker":
            # Add speaker line
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.line_spacing = 1.15
            
            # Add speaker name in bold
            if speaker:
                speaker_run = para.add_run(f"{speaker}: ")
                speaker_run.font.name = 'Calibri'
                speaker_run.font.size = Pt(12)
                speaker_run.bold = True
                speaker_run.italic = False
                content = content[len(speaker)+1:].strip() if content.startswith(speaker) else content
            
            # Add content with emphasis
            add_formatted_text(content, emphasis, para)
            
            para.paragraph_format.space_after = Pt(6)
        
        elif seg_type == "music":
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.line_spacing = 1.15
            
            music_run = para.add_run(content)
            music_run.font.name = 'Calibri'
            music_run.font.size = Pt(12)
            music_run.italic = True
            music_run.bold = False
            
            para.paragraph_format.space_after = Pt(3)
        
        elif seg_type == "scripture":
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.line_spacing = 1.15
            para.paragraph_format.left_indent = Inches(0.5)
            
            scripture_run = para.add_run(content)
            scripture_run.font.name = 'Calibri'
            scripture_run.font.size = Pt(12)
            scripture_run.bold = True
            scripture_run.italic = False
            
            para.paragraph_format.space_after = Pt(6)
        
        else:  # narration
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.line_spacing = 1.15
            para.paragraph_format.first_line_indent = Inches(0.25)
            
            add_formatted_text(content, emphasis, para)
            
            para.paragraph_format.space_after = Pt(6)
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.read()


def add_formatted_text(text: str, emphasis_list: list, para):
    """Add text to paragraph with emphasis formatting"""
    if not emphasis_list:
        # Simple text without emphasis
        run = para.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.bold = False
        run.italic = False
        return
    
    # Apply emphasis based on list
    last_pos = 0
    for emph in emphasis_list:
        emph_text = emph.get("text", "")
        style = emph.get("style", "")
        
        # Add text before emphasis
        if emph_text in text and text.index(emph_text) > last_pos:
            before = text[last_pos:text.index(emph_text)]
            if before:
                run = para.add_run(before)
                run.font.name = 'Calibri'
                run.font.size = Pt(12)
                run.bold = False
                run.italic = False
            
            # Add emphasized text
            emph_run = para.add_run(emph_text)
            emph_run.font.name = 'Calibri'
            emph_run.font.size = Pt(12)
            
            if style == "bold":
                emph_run.bold = True
                emph_run.italic = False
            elif style in ["italic", "italic_quote"]:
                emph_run.italic = True
                emph_run.bold = False
            else:
                emph_run.bold = False
                emph_run.italic = False
            
            last_pos = text.index(emph_text) + len(emph_text)
    
    # Add remaining text
    if last_pos < len(text):
        remaining = text[last_pos:]
        if remaining:
            run = para.add_run(remaining)
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.bold = False
            run.italic = False


def format_transcript(text: str, title: str = None) -> bytes:
    """Main entry point for formatting"""
    return format_transcript_with_claude(text)
