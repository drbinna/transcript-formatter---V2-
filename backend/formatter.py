from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import os
from anthropic import Anthropic
import time
import re
from dotenv import load_dotenv

# Load environment variables
load_dotenv()


def add_bottom_border(paragraph):
    """Add a bottom border to a paragraph"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    
    pPr.append(pBdr)


def get_template_path():
    """Get the template file path"""
    current_dir = os.path.dirname(os.path.abspath(__file__))
    root_dir = os.path.dirname(current_dir)
    
    # Try multiple possible locations (including in backend/ itself)
    possible_paths = [
        os.path.join(current_dir, 'sample_formatted.docx'),  # In backend/ directory
        os.path.join(current_dir, '../templates/sample_formatted.docx'),
        os.path.join(current_dir, 'templates/sample_formatted.docx'),
        os.path.join(root_dir, 'templates/sample_formatted.docx'),
        os.path.join(current_dir, '../../templates/sample_formatted.docx'),
        'templates/sample_formatted.docx',
        os.path.abspath('templates/sample_formatted.docx'),
        os.path.abspath('../templates/sample_formatted.docx'),
    ]
    
    for path in possible_paths:
        if os.path.exists(path):
            print(f"Found template at: {path}")
            return path
    
    # Debug: print all attempted paths
    print(f"Current dir: {current_dir}")
    print(f"Root dir: {root_dir}")
    print(f"Attempted paths:")
    for path in possible_paths:
        print(f"  - {path} (exists: {os.path.exists(path)})")
    
    raise FileNotFoundError("Template file not found. Checked: " + ", ".join(possible_paths))


def get_claude_client():
    """Initialize and return Claude client"""
    api_key = os.getenv('ANTHROPIC_API_KEY')
    if not api_key:
        raise ValueError("ANTHROPIC_API_KEY not found in environment")
    # Add sane defaults for hosted environments (Render) where transient network issues can occur
    return Anthropic(api_key=api_key, timeout=60, max_retries=3)


def split_text_into_chunks(text: str, max_chunk_size: int = 3000) -> list[str]:
    """
    Split transcript into smaller chunks to avoid token limits.
    Tries to split at natural boundaries (speaker changes, sentences).
    """
    if len(text) <= max_chunk_size:
        return [text]
    
    chunks = []
    current_pos = 0
    
    while current_pos < len(text):
        # Try to find a good splitting point
        chunk_end = min(current_pos + max_chunk_size, len(text))
        
        if chunk_end >= len(text):
            # Last chunk
            chunks.append(text[current_pos:])
            break
        
        # Try to find a speaker boundary
        next_speaker = text.find(':', chunk_end - 500, chunk_end)
        if next_speaker > current_pos:
            chunk_end = next_speaker + 50  # Include the speaker name
        
        chunks.append(text[current_pos:chunk_end])
        current_pos = chunk_end
    
    return chunks


def pseudo_segment_chunk(raw_text: str) -> list[dict]:
    """Heuristically split a raw chunk into segments resembling the Claude schema.
    This preserves document structure when AI calls fail.
    """
    segments: list[dict] = []
    lines = [ln.strip() for ln in raw_text.splitlines() if ln.strip()]

    scripture_pattern = re.compile(r"\b([1-3]?\s?[A-Z][a-zA-Z]+)\s+\d+:\d+(?:-\d+)?\b")
    speaker_pattern = re.compile(r"^([A-Z][A-Za-z .'-]{1,40}):\s*(.*)")

    buffer: list[str] = []
    buffer_type: str | None = None
    buffer_speaker: str | None = None

    def flush_buffer():
        nonlocal buffer, buffer_type, buffer_speaker
        if not buffer:
            return
        content = " ".join(buffer).strip()
        if not content:
            buffer = []
            buffer_type = None
            buffer_speaker = None
            return
        if buffer_type == "speaker":
            segments.append({
                "type": "speaker",
                "speaker": buffer_speaker,
                "content": content,
                "emphasis": []
            })
        elif buffer_type == "scripture":
            segments.append({
                "type": "scripture",
                "speaker": None,
                "content": content,
                "emphasis": []
            })
        elif buffer_type == "music":
            segments.append({
                "type": "music",
                "speaker": None,
                "content": content,
                "emphasis": []
            })
        else:
            segments.append({
                "type": "narration",
                "speaker": None,
                "content": content,
                "emphasis": []
            })
        buffer = []
        buffer_type = None
        buffer_speaker = None

    for ln in lines:
        # Music lines
        if '♪' in ln:
            flush_buffer()
            buffer_type = "music"
            buffer.append(ln)
            flush_buffer()
            continue

        # Scripture references
        if scripture_pattern.search(ln):
            flush_buffer()
            buffer_type = "scripture"
            buffer.append(ln)
            flush_buffer()
            continue

        # Speaker lines: Name: content
        m = speaker_pattern.match(ln)
        if m:
            flush_buffer()
            buffer_type = "speaker"
            buffer_speaker = m.group(1)
            speaker_rest = m.group(2)
            if speaker_rest:
                buffer.append(speaker_rest)
            else:
                buffer.append("")
            # do not flush immediately; allow following lines to join until next marker
            continue

        # Default narration; keep accumulating until a marker is found
        if buffer_type in (None, "narration"):
            buffer_type = "narration"
            buffer.append(ln)
        else:
            # If we were in speaker/scripture/music and encountered plain text, decide:
            # continue speaker block for coherence
            if buffer_type == "speaker":
                buffer.append(ln)
            else:
                flush_buffer()
                buffer_type = "narration"
                buffer.append(ln)

    flush_buffer()

    # Fallback single narration if segmentation failed
    if not segments:
        segments.append({"type": "narration", "speaker": None, "content": raw_text, "emphasis": []})
    return segments

def format_transcript_with_claude(text: str) -> bytes:
    """
    Use Claude AI to format the transcript according to professional standards.
    """
    client = get_claude_client()
    
    # Open the template
    template_path = get_template_path()
    doc = Document(template_path)
    
    # Prepare the comprehensive formatting instructions
    system_prompt = """You are a professional transcript formatting assistant specializing in converting raw TV show transcripts into polished, publication-ready documents with excellent readability and logical flow.

Your task is to transform unformatted transcript text into well-structured, readable content that will be formatted into a professional Word document.

Return the formatted content as JSON with this structure:
{
  "title": "Episode Title",
  "segments": [
    {
      "type": "speaker|music|scripture|narration",
      "speaker": "Speaker name or null",
      "content": "Concise, readable content (2-3 sentences max per segment to avoid truncation)",
      "emphasis": [{"text": "example", "style": "bold|italic"}]
    }
  ]
}

IMPORTANT: Keep content concise (max 200 characters per segment) to avoid token limits. Split long speeches into multiple segments if needed.

CRITICAL READABILITY REQUIREMENTS:
1. SENTENCE STRUCTURE: Break long, complex sentences into shorter, more readable sentences
2. PARAGRAPH LENGTH: Keep paragraphs to 3-5 sentences for better readability
3. LOGICAL FLOW: Ensure each paragraph has a clear topic and flows naturally to the next
4. TRANSITIONS: Add natural transitions between topics and speakers
5. CLARITY: Remove run-on sentences and fix awkward phrasing
6. PUNCTUATION: Add proper punctuation where missing
7. CAPITALIZATION: Correct improper capitalization for professional appearance
8. GRAMMAR: Fix obvious grammatical errors while PRESERVING each speaker's authentic voice and natural speaking style
   - Maintain conversational tone and informal expressions when appropriate
   - Don't over-formalize spoken language
   - Keep regional dialects, idioms, and personal speaking patterns intact
   - Only fix errors that would confuse meaning or look unprofessional

FORMATTING RULES:
- Speaker names: Bold with colon (Dr. Billy Wilson:)
- Show names: Italic (World Impact)
- Websites: Bold (worldimpact.tv)
- Organizations: Bold (ORU, Oral Roberts University)
- Scripture references: Bold (1 John 2:18, 2 Timothy 3:1-5)
- Song titles: Italic with quotes ("Give Me Jesus")
- Music/lyrics: Preserve ♪ symbols, italic text
- Quoted scripture: Italic with quotes ("But mark this...")

Transform the raw transcript to be PROFESSIONALLY READABLE with clear logical flow, proper sentence structure, polished language, and correct grammar. Maintain the original meaning, tone, and EACH SPEAKER'S AUTHENTIC VOICE - preserve their natural speaking style, conversational patterns, and personal expressions. Only fix errors that would hinder readability or professionalism.

IMPORTANT: DO NOT include standalone musical symbol segments (like "♪♪♪ ♪♪♪ ♪♪♪" by itself) that appear right after the title. Skip these noise segments to keep the transcript clean. Only include musical segments if they contain actual lyrics or meaningful musical content."""
    
    # Split transcript into chunks if needed (slightly smaller to reduce per-call output)
    chunks = split_text_into_chunks(text)
    print(f"Processing transcript in {len(chunks)} chunk(s)")
    
    all_segments = []
    title = None
    
    # Token management settings (stay under 4000 output tokens/minute)
    MAX_TOKENS_PER_CALL = 1200
    TPM_BUDGET = 3500  # keep headroom below 4000
    tpm_window_start = time.time()
    tokens_used_in_window = 0
    
    for i, chunk in enumerate(chunks):
        user_message = f"""Format this raw transcript according to professional standards:\n\n{chunk}"""
        
        # Call Claude with optimized parameters and rate-limit awareness
        # Use non-streaming to avoid connection issues
        # Simple per-minute pacing based on previous outputs
        now = time.time()
        elapsed = now - tpm_window_start
        if elapsed >= 60:
            tpm_window_start = now
            tokens_used_in_window = 0
        if tokens_used_in_window + MAX_TOKENS_PER_CALL > TPM_BUDGET:
            sleep_for = 60 - elapsed if elapsed < 60 else 0
            if sleep_for > 0:
                print(f"TPM budget nearly exhausted, sleeping {sleep_for:.1f}s to reset window...")
                time.sleep(sleep_for)
                tpm_window_start = time.time()
                tokens_used_in_window = 0

        # small jitter to avoid bursty requests in hosted envs
        try:
            time.sleep(0.5 + (hash(chunk) % 100) / 200.0)  # 0.5s to ~1.0s
        except Exception:
            pass

        attempts = 0
        backoff = 2
        while True:
            attempts += 1
            try:
                response = client.messages.create(
                    model="claude-sonnet-4-5-20250929",
                    max_tokens=MAX_TOKENS_PER_CALL,
                    temperature=0.2,
                    system=system_prompt,
                    messages=[
                        {
                            "role": "user",
                            "content": user_message
                        }
                    ],
                    stream=False
                )
                # Track usage if available
                try:
                    output_tokens = getattr(getattr(response, "usage", None), "output_tokens", None)
                    if isinstance(output_tokens, int):
                        tokens_used_in_window += output_tokens
                    else:
                        tokens_used_in_window += MAX_TOKENS_PER_CALL // 2  # conservative estimate
                except Exception:
                    tokens_used_in_window += MAX_TOKENS_PER_CALL // 2
                full_response = response.content[0].text if response.content else ""
                break
            except Exception as e:
                msg = str(e)
                # Retry on rate limit errors
                if "rate_limit" in msg or "429" in msg:
                    if attempts >= 3:
                        print(f"Rate limit after retries on chunk {i+1}: {e}. Using pseudo-segmentation fallback.")
                        fallback = pseudo_segment_chunk(chunk)
                        all_segments.extend(fallback)
                        full_response = None
                        break
                    # wait a bit longer before retrying
                    print(f"Rate limited, retrying in {backoff}s (attempt {attempts})...")
                    time.sleep(backoff)
                    backoff *= 2
                    continue
                # Retry on transient connection/server errors
                if "Connection error" in msg or "api_error" in msg or "500" in msg:
                    if attempts >= 3:
                        print(f"Network/server error after retries on chunk {i+1}: {e}. Using pseudo-segmentation fallback.")
                        fallback = pseudo_segment_chunk(chunk)
                        all_segments.extend(fallback)
                        full_response = None
                        break
                    print(f"Connection/server error, retrying in {backoff}s (attempt {attempts})...")
                    time.sleep(backoff)
                    backoff *= 2
                    continue
                # Non-retryable error => fallback
                print(f"Error calling Claude API on chunk {i+1}: {e}. Using pseudo-segmentation fallback.")
                fallback = pseudo_segment_chunk(chunk)
                all_segments.extend(fallback)
                full_response = None
                break
        
        if full_response:
            # Parse the JSON response
            # Extract JSON from markdown fences if present
            if '```json' in full_response:
                full_response = full_response.split('```json')[1].split('```')[0]
            elif '```' in full_response:
                full_response = full_response.split('```')[1]
            
            import json
            try:
                chunk_data = json.loads(full_response)
                print(f"Successfully parsed chunk {i+1} with {len(chunk_data.get('segments', []))} segments")
                
                # Store title from first chunk only
                if i == 0 and chunk_data.get("title"):
                    title = chunk_data["title"]
                
                # Collect segments from this chunk
                parsed_segments = chunk_data.get("segments", [])
                all_segments.extend(parsed_segments)
                
                # Coverage check: ensure we didn't lose too much content
                try:
                    combined_len = sum(len(seg.get("content", "")) for seg in parsed_segments)
                    if combined_len < 0.7 * len(chunk):
                        print(f"Low coverage for chunk {i+1} (parsed {combined_len} of {len(chunk)} chars). Appending raw chunk to preserve completeness.")
                        all_segments.append({"type": "narration", "content": chunk, "emphasis": []})
                except Exception:
                    pass
                
            except json.JSONDecodeError as e:
                print(f"JSON parsing error in chunk {i+1}: {e}")
                
                # Try to repair truncated JSON
                try:
                    last_brace = full_response.rfind('}')
                    second_last_brace = full_response.rfind('}', 0, last_brace) if last_brace > 0 else -1
                    
                    if last_brace > 0 and second_last_brace > 0:
                        repaired = full_response[:second_last_brace+1] + ']}'
                        try:
                            chunk_data = json.loads(repaired)
                            print(f"Successfully repaired chunk {i+1} JSON with {len(chunk_data.get('segments', []))} segments")
                            if i == 0 and chunk_data.get("title"):
                                title = chunk_data["title"]
                            all_segments.extend(chunk_data.get("segments", []))
                        except json.JSONDecodeError:
                            repaired = full_response[:last_brace+1] + ']}'
                            chunk_data = json.loads(repaired)
                            print(f"Successfully repaired chunk {i+1} JSON with {len(chunk_data.get('segments', []))} segments")
                            if i == 0 and chunk_data.get("title"):
                                title = chunk_data["title"]
                            all_segments.extend(chunk_data.get("segments", []))
                except Exception as repair_err:
                    print(f"JSON repair failed for chunk {i+1}: {repair_err}. Using pseudo-segmentation fallback.")
                    fallback = pseudo_segment_chunk(chunk)
                    all_segments.extend(fallback)
        else:
            # No model response (e.g., after retries) -> pseudo segment
            fallback = pseudo_segment_chunk(chunk)
            all_segments.extend(fallback)
    
    # Combine all segments into one structure
    data = {
        "title": title or "World Impact Transcript",
        "segments": all_segments
    }
    
    print(f"Total segments after combining: {len(all_segments)}")
    
    # Apply formatting to document
    if data.get("title"):
        title_para = doc.add_paragraph(data["title"])
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(18)
            run.bold = True
        title_para.paragraph_format.space_before = Pt(12)
        title_para.paragraph_format.space_after = Pt(18)
        add_bottom_border(title_para)
    
    # Process segments
    prev_seg_type = None
    
    for i, segment in enumerate(data.get("segments", [])):
        seg_type = segment.get("type", "narration")
        content = segment.get("content", "")
        speaker = segment.get("speaker")
        emphasis = segment.get("emphasis", [])
        
        # Filter out pure musical symbol noise (after title)
        if i <= 2 and content.strip() in ['♪♪♪', '♪♪', '♪♪♪ ♪♪♪', '♪♪♪ ♪♪♪ ♪♪♪', '♪ ♪ ♪']:
            continue  # Skip pure musical symbol segments right after title
        
        # Add extra break when transitioning between different content types
        if prev_seg_type and prev_seg_type != seg_type:
            # Add extra spacing for major transitions
            blank_para = doc.add_paragraph()
            blank_para.paragraph_format.space_after = Pt(3)
        
        # Track previous values for transitions
        prev_seg_type = seg_type
        
        if seg_type == "speaker":
            # Add speaker line
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.line_spacing = 1.2  # Increased for better readability
            para.paragraph_format.space_before = Pt(3)  # Add space before speaker for clarity
            
            # Add speaker name in bold
            if speaker:
                speaker_run = para.add_run(f"{speaker}: ")
                speaker_run.font.name = 'Calibri'
                speaker_run.font.size = Pt(12)
                speaker_run.bold = True
                speaker_run.italic = False
                content = content[len(speaker)+1:].strip() if content.startswith(speaker) else content
            
            # Add content with emphasis
            add_formatted_text(content, emphasis, para)
            
            para.paragraph_format.space_after = Pt(9)  # Increased for better readability
        
        elif seg_type == "music":
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.line_spacing = 1.2  # Increased for better readability
            
            music_run = para.add_run(content)
            music_run.font.name = 'Calibri'
            music_run.font.size = Pt(12)
            music_run.italic = True
            music_run.bold = False
            
            para.paragraph_format.space_after = Pt(3)
        
        elif seg_type == "scripture":
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.line_spacing = 1.2  # Increased for better readability
            para.paragraph_format.left_indent = Inches(0.5)
            
            scripture_run = para.add_run(content)
            scripture_run.font.name = 'Calibri'
            scripture_run.font.size = Pt(12)
            scripture_run.bold = True
            scripture_run.italic = False
            
            para.paragraph_format.space_after = Pt(9)  # Increased for better readability
        
        else:  # narration
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.line_spacing = 1.2  # Increased for better readability
            para.paragraph_format.first_line_indent = Inches(0.25)
            
            add_formatted_text(content, emphasis, para)
            
            para.paragraph_format.space_after = Pt(9)  # Increased for better readability
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.read()


def add_formatted_text(text: str, emphasis_list: list, para):
    """Add text to paragraph with emphasis formatting using regex for reliability"""
    import re
    
    if not emphasis_list:
        # Simple text without emphasis
        run = para.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.bold = False
        run.italic = False
        return
    
    # Sort emphasis by occurrence in text, avoiding duplicates
    emph_positions = []
    seen_positions = set()
    for emph in emphasis_list:
        emph_text = emph.get("text", "")
        style = emph.get("style", "")
        # Find all occurrences
        for match in re.finditer(re.escape(emph_text), text):
            pos_key = (match.start(), match.end())
            if pos_key not in seen_positions:
                emph_positions.append((match.start(), match.end(), emph_text, style))
                seen_positions.add(pos_key)
    
    if not emph_positions:
        # No matches found, just add the text
        run = para.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.bold = False
        run.italic = False
        return
    
    # Sort by position
    emph_positions.sort()
    
    # Apply emphasis (handle overlapping spans by taking the maximum)
    last_pos = 0
    for i, (start, end, emph_text, style) in enumerate(emph_positions):
        # Skip if this span is already covered
        if start < last_pos:
            continue
            
        # Add text before this emphasis
        if start > last_pos:
            before = text[last_pos:start]
            if before:
                run = para.add_run(before)
                run.font.name = 'Calibri'
                run.font.size = Pt(12)
                run.bold = False
                run.italic = False
        
        # Add emphasized text
        emph_run = para.add_run(emph_text)
        emph_run.font.name = 'Calibri'
        emph_run.font.size = Pt(12)
        
        if style == "bold":
            emph_run.bold = True
            emph_run.italic = False
        elif style in ["italic", "italic_quote"]:
            emph_run.italic = True
            emph_run.bold = False
        else:
            emph_run.bold = False
            emph_run.italic = False
        
        last_pos = end
    
    # Add remaining text
    if last_pos < len(text):
        remaining = text[last_pos:]
        if remaining:
            run = para.add_run(remaining)
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.bold = False
            run.italic = False


def format_basic_transcript(text: str) -> bytes:
    """
    Fallback formatting function when AI fails.
    Returns a basic formatted document.
    """
    try:
        template_path = get_template_path()
        doc = Document(template_path)
        
        # Add a title paragraph
        title_para = doc.add_paragraph("Transcript")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(18)
            run.bold = True
        title_para.paragraph_format.space_after = Pt(18)
        add_bottom_border(title_para)
        
        # Add content
        content_para = doc.add_paragraph(text)
        content_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        content_para.paragraph_format.line_spacing = 1.2
        content_para.paragraph_format.first_line_indent = Inches(0.25)
        content_para.paragraph_format.space_after = Pt(9)
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.read()
    except Exception as e:
        print(f"Error in format_basic_transcript: {e}")
        raise


def format_transcript(text: str, title: str = None) -> bytes:
    """Main entry point for formatting"""
    try:
        return format_transcript_with_claude(text)
    except Exception as e:
        print(f"Error in format_transcript: {e}")
        # Fallback to basic formatting
        return format_basic_transcript(text)
