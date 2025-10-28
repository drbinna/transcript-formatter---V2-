"""
Professional Transcript Formatter
Following ORU formatting standards for sermon transcripts
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING
import io
import re

def add_bottom_border(paragraph):
    """Add a bottom border to a paragraph"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    # Create a bottom border element
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')      # Line thickness (1/2 pt)
    bottom.set(qn('w:space'), '1')   # Padding
    bottom.set(qn('w:color'), '000000')  # Black
    pBdr.append(bottom)
    
    pPr.append(pBdr)

def format_transcript(text: str) -> bytes:
    """
    Format a raw transcript into a professional ORU-formatted .docx document.
    
    Args:
        text: Raw transcript content as a string
        
    Returns:
        bytes: The formatted .docx file as bytes
    """
    
    lines = text.split('\n')
    
    # Create a new document
    doc = Document()
    
    # Set page size and margins
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # Track state
    title_added = False
    prev_was_speaker = False
    
    for line in lines:
        line = line.strip()
        
        # Skip empty lines initially, then add appropriate spacing
        if not line:
            if title_added:
                doc.add_paragraph()  # Add blank line
            continue
        
        # Detect title (first meaningful line)
        if not title_added:
            title_para = doc.add_paragraph(line)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Format title
            for run in title_para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(18)
                run.bold = True
            
            # Set spacing: 12 pt before, 18 pt after
            title_para.paragraph_format.space_before = Pt(12)
            title_para.paragraph_format.space_after = Pt(18)
            
            # Add bottom border
            add_bottom_border(title_para)
            
            title_added = True
            prev_was_speaker = False
            continue
        
        # Detect speaker lines (name ending with colon)
        if re.match(r'^[A-Z][a-zA-Z\s\.]+:$', line) or ':' in line and len(line) < 100:
            # Check if it's a speaker pattern (name: followed by text)
            parts = line.split(':', 1)
            if len(parts) > 1:
                # Split speaker name and content
                speaker_name = parts[0] + ':'
                content = parts[1].strip()
                
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                para.paragraph_format.line_spacing = 1.15
                
                # Add speaker name in bold
                speaker_run = para.add_run(speaker_name)
                speaker_run.font.name = 'Calibri'
                speaker_run.font.size = Pt(12)
                speaker_run.bold = True
                
                # Add content in normal
                if content:
                    content_run = para.add_run(' ' + content)
                    content_run.font.name = 'Calibri'
                    content_run.font.size = Pt(12)
                    content_run.bold = False
                
                para.paragraph_format.space_after = Pt(6)
                prev_was_speaker = True
            else:
                # Just speaker name
                para = doc.add_paragraph(line)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                para.paragraph_format.line_spacing = 1.15
                
                for run in para.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(12)
                    run.bold = True
                
                para.paragraph_format.space_after = Pt(6)
                prev_was_speaker = True
            continue
        
        # Detect music lines
        if '♪' in line:
            para = doc.add_paragraph(line)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.line_spacing = 1.15
            
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(12)
                run.italic = True
            
            para.paragraph_format.space_after = Pt(3)
            prev_was_speaker = False
            continue
        
        # Detect scripture references
        if re.match(r'^\d+\s+[A-Z][a-zA-Z]+\s+\d+:\d+', line) or re.search(r'\d+\s+[A-Z][a-zA-Z]+\s+\d+:\d+', line):
            para = doc.add_paragraph(line)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.line_spacing = 1.15
            para.paragraph_format.left_indent = Inches(0.5)
            
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(12)
                run.italic = True
            
            para.paragraph_format.space_after = Pt(6)
            prev_was_speaker = False
            continue
        
        # Regular narration/descriptive text
        para = doc.add_paragraph(line)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para.paragraph_format.line_spacing = 1.15
        para.paragraph_format.first_line_indent = Inches(0.25)
        
        for run in para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.bold = False
        
        para.paragraph_format.space_after = Pt(6)
        prev_was_speaker = False
    
    # Add footer with ORU branding
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Oral Roberts University Presents: World Impact with Dr. Billy Wilson"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if footer_para.runs:
        footer_run = footer_para.runs[0]
        footer_run.font.size = Pt(10)
        footer_run.font.name = 'Calibri'
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.read()

