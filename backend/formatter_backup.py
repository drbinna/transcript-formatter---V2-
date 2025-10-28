from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import io
import os
import json
from anthropic import Anthropic

# Get Claude client - will be initialized when API key is loaded
claude_client = None

def get_claude_client():
    """Get or initialize Claude client with API key from environment"""
    global claude_client
    if claude_client is None:
        api_key = os.getenv("ANTHROPIC_API_KEY")
        if not api_key:
            raise ValueError("ANTHROPIC_API_KEY not found in environment variables")
        claude_client = Anthropic(api_key=api_key)
    return claude_client

def read_template_context(template_path: str = None) -> str:
    """Read the template document to use as context for Claude"""
    if template_path is None:
        # Try multiple possible paths
        import os
        current_dir = os.path.dirname(os.path.abspath(__file__))
        possible_paths = [
            os.path.join(current_dir, "..", "templates", "sample_formatted.docx"),
            os.path.join(current_dir, "templates", "sample_formatted.docx"),
            "templates/sample_formatted.docx",
        ]
        
        for path in possible_paths:
            if os.path.exists(path):
                template_path = path
                break
        
        if template_path is None or not os.path.exists(template_path):
            print(f"Warning: Could not find template at any of: {possible_paths}")
            return ""
    
    try:
        doc = Document(template_path)
        # Extract first 20 paragraphs as context
        context = []
        for para in doc.paragraphs[:20]:
            if para.text.strip():
                context.append(para.text)
        return "\n".join(context)
    except Exception as e:
        print(f"Warning: Could not read template: {e}")
        return ""

def add_bottom_border(paragraph):
    """Add a bottom border to a paragraph"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    # Create a bottom border element
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')      # Line thickness
    bottom.set(qn('w:space'), '1')   # Padding
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    
    pPr.append(pBdr)

def format_transcript(text: str, title: str = None) -> bytes:
    """
    Use Claude AI to intelligently parse and format transcript based on sample template.
    
    Args:
        text: Raw transcript content as a string
        title: Optional title for the document
        
    Returns:
        bytes: The formatted .docx document as binary data
    """
    # Read the template to understand the desired formatting
    template_text = read_template_context()
    
    # Use Claude to analyze and format the transcript
    client = get_claude_client()
    response = client.messages.create(
        model="claude-sonnet-4-5-20250929",  # Claude Sonnet 4.5
        max_tokens=8192,  # Max tokens for Sonnet 4.5
        stream=True,  # Enable streaming for reliable response handling
        system="""You are a document formatting assistant. Your task is to analyze a raw transcript and format it according to the example template provided.

The formatting rules based on the template:
1. Title: Centered, bold, 20pt (Gotham font)
2. Speaker names (ending with ':'): Bold, 12pt Times New Roman
3. Scripture references (e.g., '1 John 2:18', '2 Timothy 3:1-5'): Bold
4. Text in double quotes (show names, song titles, scripture quotes): Italic
5. Music lines (containing ♪): Preserve as-is
6. Regular text: 12pt Times New Roman

Format the transcript using this JSON structure:
{
  "title": "Document title",
  "paragraphs": [
    {"text": "Dr. Billy Wilson:", "bold": true},
    {"text": "Welcome to the show...", "bold": false}
  ]
}

IMPORTANT FORMATTING RULES:
- Speaker names (ending with colon ':'): Set "bold": true
- "World Impact" brand name: Set "italic": true  
- Song titles in quotes (e.g. "Give Me Jesus"): Set "italic": true
- Direct quotes/speech: Set "italic": true
- Regular text: Set "bold": false, "italic": false
- Brand names like "worldimpact.tv" and "ORU" should remain regular text (NOT bold)
- Music lines with ♪: Preserve exactly as-is  
- CRITICAL: Ensure complete sentences. Do NOT create text segments where punctuation appears without proper context
- Each text segment should be a complete thought or sentence fragment that maintains proper grammatical structure
- Avoid splitting text in ways that would cause punctuation to be orphaned at the start of lines

Return ONLY valid JSON, no markdown formatting.""",
        messages=[
            {
                "role": "user",
                "content": f"""Here is the sample formatting template:

{template_text}

Now format this transcript:

{text}

Provide the formatted result as JSON following the structure above."""
            }
        ]
    )
    
    # Parse Claude's response (handle streaming)
    response_text = ""
    for event in response:
        if event.type == "content_block_delta":
            response_text += event.delta.text
        elif event.type == "content_block_start":
            response_text += event.content_block.text
    
    # Strip markdown code blocks if present
    if response_text.startswith("```json"):
        response_text = response_text[7:]  # Remove ```json
    if response_text.startswith("```"):
        response_text = response_text[3:]  # Remove ```
    if response_text.endswith("```"):
        response_text = response_text[:-3]  # Remove trailing ```
    
    response_text = response_text.strip()
    
    # Try to parse JSON with better error handling
    # Save response for debugging
    with open('/tmp/claude_response.txt', 'w') as f:
        f.write(response_text)
    
    try:
        formatted_json = json.loads(response_text)
    except json.JSONDecodeError as e:
        print(f"JSON Parse Error at position {e.pos}: {e.msg}")
        print(f"Response length: {len(response_text)}")
        
        # Try to fix incomplete JSON by closing brackets/braces
        import re
        response_text = response_text.strip()
        
        # Count unmatched brackets
        open_braces = response_text.count('{') - response_text.count('}')
        open_brackets = response_text.count('[') - response_text.count(']')
        
        # Fix incomplete strings
        if response_text[-1] != '"':
            # Find the last opening quote and close it
            last_quote = response_text.rfind('"')
            if last_quote > 0:
                response_text = response_text[:last_quote+1] + '"'
        
        # Close open brackets
        response_text += '"}]' * (open_brackets)
        response_text += '}' * (open_braces)
        
        try:
            formatted_json = json.loads(response_text)
            print("Successfully fixed and parsed incomplete JSON")
        except json.JSONDecodeError:
            # Last resort: create minimal valid JSON from what we have
            import re
            title_match = re.search(r'"title"\s*:\s*"([^"]+)"', response_text)
            paragraphs_match = re.findall(r'"text"\s*:\s*"([^"]+)"', response_text)
            
            formatted_json = {
                "title": title_match.group(1) if title_match else "Living in the Last Days",
                "paragraphs": [{"text": p} for p in paragraphs_match]
            }
            print("Used regex-based extraction as last resort")
    
    # Create Word document
    doc = Document()
    
    # Add title with bottom border
    title_para = doc.add_paragraph(formatted_json["title"])
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.font.size = Pt(20)
    title_run.bold = True
    title_run.font.name = 'Gotham'
    
    # Add bottom border to title
    add_bottom_border(title_para)
    
    doc.add_paragraph()
    
    # Add formatted paragraphs with professional spacing
    for i, para_data in enumerate(formatted_json["paragraphs"]):
        para = doc.add_paragraph()
        
        # Handle new structure with direct text, bold, italic fields
        text = para_data.get("text", "")
        
        # Ensure proper text formatting - add non-breaking spaces to prevent punctuation at line starts
        # Replace spaces before punctuation with non-breaking spaces
        import re
        # Use actual non-breaking space character instead of escape
        nbsp = '\u00a0'
        text = re.sub(r' ([,.!?:;])', nbsp + r'\1', text)
        
        run = para.add_run(text)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.bold = para_data.get("bold", False)
        run.italic = para_data.get("italic", False)
        
        # Professional paragraph spacing
        para.paragraph_format.space_after = Pt(3)
        
        # Add spacing after speaker names and music for clarity
        text_lower = text.lower()
        if "♪" in text or (":" in text and len(text) < 80 and text.endswith(":")):
            # Add minimal spacing after speaker lines
            para.paragraph_format.space_after = Pt(6)
    
    # Add ORU branding footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Oral Roberts University Presents: World Impact with Dr. Billy Wilson"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.runs[0]
    footer_run.font.size = Pt(10)
    footer_run.font.name = 'Times New Roman'
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.read()

